# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Vagas.me (Fidelidade ao Aplicativo)
DARK   = (13,  14,  18)   # #0D0E12
PURPLE = (139, 38,  253)  # #8B26FD (Electric Purple)
AMBER  = (255, 159, 10)   # #FF9F0A (Orange/Amber)
GREEN  = (16,  185, 129)  # #10B981 (Emerald Success)
GRAY   = (160, 174, 192)  # #A0AEC0 (Muted Blue/Gray)
WHITE  = (255, 255, 255)

def shade(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str)
    tcPr.append(shd)

def tbl(doc, headers, rows, header_color='8B26FD', alt_color='F5F3FF'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        shade(c, header_color)
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        r = t.rows[ri+1]
        for ci, val in enumerate(row):
            c = r.cells[ci]
            c.text = str(val)
            run = c.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if ri % 2 == 0:
                shade(c, alt_color)
    return t

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(*PURPLE)
    return p

def h2(doc, text, color=PURPLE):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(*color)
    return p

def h3(doc, text, color=PURPLE):
    p = doc.add_heading(text, level=3)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(*color)
    return p

def par(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(size)
    return p

def divider(doc):
    doc.add_paragraph('─' * 85)

# ── BUILD DOCUMENT ──────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.2); s.right_margin = Inches(1.2)
doc.styles['Normal'].font.name = 'Plus Jakarta Sans'
doc.styles['Normal'].font.size = Pt(11)

# ── CAPA ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Vagas.me')
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(*PURPLE)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run('PLANEJAMENTO ESTRATÉGICO & OKRs')
r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(*DARK)

s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = s2.add_run('"O seu trabalho perto de você."')
r3.italic = True; r3.font.size = Pt(13); r3.font.color.rgb = RGBColor(*GRAY)

doc.add_paragraph(); doc.add_paragraph()
s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = s3.add_run('Equipe Fundadora Vagas.me  |  Co-piloto: Antigravity  |  Julho 2026')
r4.font.size = Pt(10); r4.font.color.rgb = RGBColor(*GRAY)

doc.add_paragraph('=' * 85)
doc.add_page_break()

# ── ÍNDICE VISUAL ───────────────────────────────────────────────────────────
h1(doc, 'ESTRUTURA DO DOCUMENTO')
secs = [
    ('1', 'Missão, Visão e Valores (Foco no Curto Deslocamento)'),
    ('2', 'Estrutura e Divisão de Quotas (Divisão de Papéis)'),
    ('3', 'OKRs — Objetivos e Resultados-Chave (Q3 + Q4 2026)'),
    ('4', 'KPIs — Dashboard de Métricas Operacionais'),
    ('5', 'Metas Anuais 2026 e 2027'),
    ('6', 'Planejamento Financeiro & Break-even'),
    ('7', 'Roadmap de Produto (Web App MVP -> App Autônomos)'),
    ('8', 'Rituais de Gestão da Sociedade'),
    ('9', 'Regras da Casa & Entregas'),
]
for num, title in secs:
    p = doc.add_paragraph()
    r_n = p.add_run(f'  {num}.  '); r_n.bold = True; r_n.font.color.rgb = RGBColor(*PURPLE)
    r_t = p.add_run(title); r_t.font.size = Pt(11)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — MISSÃO, VISÃO E VALORES
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 1 — MISSÃO, VISÃO E VALORES')
par(doc, 'A essência do Vagas.me é a simplicidade e a utilidade real. Rejeitamos a complexidade corporativa e focamos em vagas CLT locais.', italic=True)
doc.add_paragraph()

h2(doc, '🎯 Missão')
par(doc, '"Aproximar o trabalhador CLT das oportunidades de emprego em sua própria região por meio de um mapa interativo, eliminando fricção e reduzindo o tempo de deslocamento urbano."', italic=True, size=13, color=PURPLE)
doc.add_paragraph()

h2(doc, '🔭 Visão')
par(doc, '"Validar o modelo geolocalizado de contratação em Itapetininga no primeiro trimestre, expandir regionalmente no segundo semestre e atrair fundos de investimento de alta escala em até 12 meses."', italic=True, size=12, color=DARK)
doc.add_paragraph()

h2(doc, '💎 Valores')
valores = [
    ('Proximidade Física', 'O foco é no deslocamento. O melhor trabalho é o trabalho perto de casa.'),
    ('Simplicidade Extrema', 'Zero feeds poluídos. Exibição direta em mapa e filtros intuitivos.'),
    ('Democracia de Acesso', 'Gerador de currículo 100% gratuito e sem cobranças abusivas de candidatos.'),
    ('Escala Orientada a Rede', 'Foco total no efeito de rede. A base de usuários ativos é o nosso maior castelo de defesa.'),
]
tbl(doc, ['Valor', 'O que significa na prática'], valores)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — ESTRUTURA E SOCIEDADE
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 2 — ESTRUTURA OPERACIONAL')
par(doc, 'A divisão inicial do projeto foca na especialização por área para acelerar a entrega do MVP.', italic=True)
doc.add_paragraph()

h2(doc, 'Organograma & Responsabilidades')
tbl(doc,
    ['Área / Papel', 'Foco Central', 'Responsabilidades principais'],
    [
        ['Área Social / Parcerias', 'Crescimento e Investimento', 'Conexão com investidores, palestra em eventos de empreendedorismo, parcerias com empresas locais (ex: Cofesa)'],
        ['Desenvolvimento Web (Site)', 'MVP do Web App e Mapa', 'Implementação do mapa interativo com Leaflet.js, API de vagas, painel Kanban e gerador de currículos'],
        ['Desenvolvimento App Mobile', 'Aplicativo Nativo Futuro', 'Estruturação da versão mobile (iOS/Android), geolocalização nativa e futuras rotas de profissionais autônomos'],
    ],
    header_color='8B26FD'
)
doc.add_paragraph()

h2(doc, 'Fases de Formalização Societária')
tbl(doc,
    ['Fase', 'Objetivo do Acordo', 'Gatilho de Execução'],
    [
        ['FASE 0 — Acordo de Parceria', 'Assinatura de Acordo de Co-fundadores (Vesting inicial)', 'Imediato (Início do MVP)'],
        ['FASE 1 — Validação Local', 'Teste do Web App em Itapetininga por 1 a 3 meses', 'Finalização da versão 1.0 do site'],
        ['FASE 2 — Abertura de Sociedade Ltda', 'Consolidação das participações e quotas societárias', 'Faturamento inicial ou aporte de capital'],
        ['FASE 3 — Rodada de Investimento', 'Apresentação do Pitch-Deck estruturado para grupos de investimento', 'Tração local comprovada com 1.000+ currículos'],
    ],
    header_color='1F202B'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — OKRs
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 3 — OKRs (Objectives & Key Results)')
doc.add_paragraph()

# Q3
h2(doc, '⚡ Q3 2026 — FASE: "FUNDAÇÃO & VALIDAÇÃO LOCAL"', color=PURPLE)
par(doc, 'Foco do trimestre: lançar o MVP Web App e iniciar o teste de mercado na cidade de Itapetininga.', italic=True)
doc.add_paragraph()

okrs_q3 = [
    ('O1', 'Lançar e validar o MVP do Web App funcional no mapa real',
     [
         ('KR1', 'Finalizar o mapa interativo funcional com Leaflet.js', 'Não', 'Sim'),
         ('KR2', 'Cadastrar e manter ativas 15 empresas parceiras de Itapetininga', '0', '15'),
         ('KR3', 'Atingir 1.000 currículos gerados gratuitamente na plataforma', '0', '1.000'),
         ('KR4', 'Alcançar 300 candidaturas geolocalizadas efetuadas no mapa', '0', '300'),
     ]),
    ('O2', 'Consolidar a estrutura de marketing de atração inicial',
     [
         ('KR1', 'Criar perfil do Instagram Vagas.me e obter 500 seguidores locais', '0', '500'),
         ('KR2', 'Realizar divulgação direta em 30 comércios locais (padarias, supermercados)', '0', '30'),
         ('KR3', 'Atingir 50 downloads diários do PDF de currículo com a marca d\'água do Vagas.me', '0', '50'),
     ]),
    ('O3', 'Estruturar o Pitch-Deck para captação de investimento',
     [
         ('KR1', 'Documentar a jornada e o custo de atração do candidato em Itapetininga', 'Não', 'Sim'),
         ('KR2', 'Apresentar o MVP e o plano estratégico para 2 mentores de investimento', '0', '2'),
     ]),
]

for obj_id, obj_title, krs in okrs_q3:
    h3(doc, f'{obj_id} — {obj_title}', color=PURPLE)
    tbl(doc,
        ['KR', 'Key Result', 'Baseline', 'Meta Q3'],
        [(kr_id, kr_desc, bl, meta) for kr_id, kr_desc, bl, meta in krs],
        header_color='8B26FD', alt_color='F5F3FF'
    )
    doc.add_paragraph()

doc.add_page_break()

# Q4
h2(doc, '🚀 Q4 2026 — FASE: "TRAÇÃO REGIONAL & INFRAESTRUTURA"', color=AMBER)
par(doc, 'Foco do trimestre: expansão regional de candidatos ativos e preparação da primeira rodada de captação.', italic=True)
doc.add_paragraph()

okrs_q4 = [
    ('O1', 'Expandir a presença do mapa para cidades da microrregião',
     [
         ('KR1', 'Cadastrar vagas e empresas em 3 cidades polo vizinhas', '1', '4'),
         ('KR2', 'Atingir 5.000 candidatos ativos na base de dados (efeito de rede)', '1.000', '5.000'),
         ('KR3', 'Alcançar 1.500 candidaturas mensais no mapa', '300', '1.500'),
     ]),
    ('O2', 'Apresentar dados do MVP para investidores de alto nível',
     [
         ('KR1', 'Apresentar a plataforma em 2 eventos de empreendedorismo', '0', '2'),
         ('KR2', 'Realizar pitch para 3 investidores ligados ao ecossistema do Marçal', '0', '3'),
     ]),
    ('O3', 'Testar e validar a primeira receita recorrente (Plus)',
     [
         ('KR1', 'Atingir 100 assinantes ativos do Plano Plus (R$ 7,90/mês)', '0', '100'),
         ('KR2', 'Obter R$ 1.000 de faturamento em anúncios de tela cheia ou banners patrocinados', 'R$ 0', 'R$ 1.000'),
     ]),
]

for obj_id, obj_title, krs in okrs_q4:
    h3(doc, f'{obj_id} — {obj_title}', color=AMBER)
    tbl(doc,
        ['KR', 'Key Result', 'Baseline', 'Meta Q4'],
        [(kr_id, kr_desc, bl, meta) for kr_id, kr_desc, bl, meta in krs],
        header_color='FF9F0A', alt_color='FFFBEB'
    )
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 — KPIs
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 4 — KPIs (Dashboard de Métricas)')
par(doc, 'Revisão semanal para monitoramento do crescimento do efeito de rede.', italic=True)
doc.add_paragraph()

h2(doc, '📱 KPIs de Usuários e Efeito de Rede')
tbl(doc,
    ['Métrica', 'O que mede', 'Frequência', 'Meta Mês 3'],
    [
        ['Currículos criados', 'Engajamento com o gerador gratuito', 'Semanal', '1.000+'],
        ['Visualizações de currículo com marca d\'água', 'Disseminação orgânica da marca Vagas.me', 'Semanal', '500+'],
        ['Acessos únicos ao mapa', 'Interesse de candidatos', 'Diário', '200+ acessos/dia'],
        ['Candidaturas efetuadas', 'Volume de matches reais', 'Mensal', '300+'],
        ['Empresas cadastradas', 'Disponibilidade de vagas no mapa', 'Mensal', '15+'],
    ],
    header_color='8B26FD', alt_color='F5F3FF'
)
doc.add_paragraph()

h2(doc, '💰 KPIs de Monetização (Simulada no MVP)')
tbl(doc,
    ['Métrica', 'O que mede', 'Meta Q3', 'Meta Q4'],
    [
        ['Visualizações de anúncios (Ads)', 'Tempo pago pelos desempregados', '10.000 visualizações', '50.000 visualizações'],
        ['Clicks em Banners patrocinados', 'Interesse em anúncios locais', '500 clicks', '2.000 clicks'],
        ['Assinantes Vagas.me Plus', 'Candidatos buscando prioridade sem ads', '10 assinantes teste', '100 assinantes ativos'],
        ['Empresas pagando por destaque', 'Volume de marcas em destaque no mapa', '2 empresas', '8 empresas'],
    ],
    header_color='1F202B', alt_color='F8F8FF'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 5 — METAS ANUAIS
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 5 — METAS ANUAIS 2026 E 2027')
doc.add_paragraph()

h2(doc, '📅 Metas 2026 (MVP & Validação Regional)')
tbl(doc,
    ['Área', 'Meta', 'Prazo'],
    [
        ['MVP do Web App', 'Finalizar mapa interativo e gerador de currículos funcionais', 'Julho/26'],
        ['Teste local', 'Rodar em Itapetininga de forma fechada', 'Julho a Setembro/26'],
        ['Audiência de Rede', 'Obter 2.000 candidatos cadastrados', 'Outubro/26'],
        ['Expansão regional', 'Lançar nas cidades vizinhas (Sorocaba, Tatuí, etc.)', 'Novembro/26'],
        ['Investidores', 'Apresentar dados consolidados para 3 investidores anjo', 'Dezembro/26'],
    ],
    header_color='8B26FD', alt_color='F5F3FF'
)
doc.add_paragraph()

h2(doc, '📅 Metas 2027 (Escala Nacional & App Nativo)')
tbl(doc,
    ['Área', 'Meta', 'Prazo'],
    [
        ['Aplicativo Nativo', 'Desenvolvimento e lançamento nas lojas (iOS/Android)', 'Março/27'],
        ['Autônomos no Mapa', 'Inclusão de geolocalização de autônomos/prestadores de serviços', 'Junho/27'],
        ['Crescimento de Base', 'Alcançar 100.000 currículos gerados a nível nacional', 'Setembro/27'],
        ['Valuation & Aporte', 'Captação de rodada de investimento institucional (Série A)', 'Dezembro/27'],
    ],
    header_color='FF9F0A', alt_color='FFFBEB'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 6 — FINANCEIRO
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 6 — PLANEJAMENTO FINANCEIRO')
par(doc, 'No início, o orçamento é estritamente baixo (Low Budget) em software, concentrando os recursos financeiros na divulgação e no marketing de escala.', italic=True)
doc.add_paragraph()

h2(doc, '💸 Custos Operacionais Mínimos (Servidor / IA / APIs)')
tbl(doc,
    ['Serviço', 'Tipo', 'Valor Mensal Estimado', 'Observação'],
    [
        ['Banco de Dados PostgreSQL (Supabase)', 'Fixo / Armazenamento', 'R$ 0,00 (Plano Free)', 'Gratuito para os primeiros 500MB de dados'],
        ['Hospedagem Frontend (Vercel)', 'Fixo / Site', 'R$ 0,00 (Plano Free)', 'Gratuito para hospedagem de SPA'],
        ['Servidor Backend (Railway/Render)', 'Fixo / Backend Python', 'R$ 35,00', 'Para execução do FastAPI/Flask'],
        ['API de Match/Descrição (Gemini API)', 'Variável (por uso)', 'R$ 15,00', 'Pagamento proporcional ao número de requisições'],
        ['API do Mapa (Mapbox/OSM)', 'Fixo / Exibição', 'R$ 0,00', 'Gratuito até o limite generoso de requisições'],
        ['TOTAL FIXO OPERACIONAL', '—', 'R$ 50,00 / mês', 'Estrutura inicial ultra enxuta'],
    ],
    header_color='1F202B', alt_color='F8F8FF'
)
doc.add_paragraph()

h2(doc, '💰 Projeção de Faturamento Estimado (Q4 2026)')
tbl(doc,
    ['Fonte de Receita', 'Preço Unitário', 'Volume Estimado', 'Projeção Mensal'],
    [
        ['Assinatura Plus', 'R$ 7,90 / mês', '100 assinantes ativos', 'R$ 790,00'],
        ['Destaques no Mapa (CNPJ)', 'R$ 49,90 / mês', '10 empresas ativas', 'R$ 499,00'],
        ['Anúncios In-App (Ads)', 'CPM (por visualizações)', '30.000 visualizações', 'R$ 200,00'],
        ['TOTAL PROJETADO', '—', '—', 'R$ 1.489,00 / mês'],
    ],
    header_color='10B981', alt_color='ECFDF5'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 7 — ROADMAP DE PRODUTO
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 7 — ROADMAP DE DESENVOLVIMENTO')
doc.add_paragraph()

roadmap = [
    ('Fase 1 (MVP)', 'Lançamento do Web App com mapa básico de Itapetininga, busca por tags de cargos, gerador de currículos gratuito em PDF imprimível e painel de controle simples.'),
    ('Fase 2 (IA)', 'Implementação do algoritmo local de match semântico no banco de dados. Integração do assistente de redação de vagas para empresas.'),
    ('Fase 3 (Plus)', 'Ativação do gateway de assinaturas Plus (PIX/Cartão), ativação de banners e anúncios em vídeo pop-up de 5 segundos no modo gratuito.'),
    ('Fase 4 (App)', 'Migração do Web App para aplicativo mobile nativo (iOS/Android). Inclusão da busca geolocalizada de profissionais autônomos (encanadores, pintores) no mapa.'),
]
tbl(doc, ['Fase de Produto', 'Entregáveis e Objetivos'], roadmap)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 8 — RITUAIS DE GESTÃO
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 8 — RITUAIS DE GESTÃO DA SOCIEDADE')
par(doc, 'A rotina de reuniões de alinhamento garante a tração constante da sociedade antes do lançamento público.', italic=True)
doc.add_paragraph()

h2(doc, '🔄 Rituais Definidos')
tbl(doc,
    ['Ritual', 'Frequência', 'Participantes', 'Objetivo Central'],
    [
        ['Alinhamento Semanal (Weekly)', 'Toda Segunda-feira', 'Co-fundadores', 'Ajustar o backlog de desenvolvimento, revisar feedbacks do MVP e planejar ações comerciais da semana.'],
        ['Revisão de Caixa e Parcerias', 'Mensal (1ª semana)', 'Co-fundadores', 'Revisar custos de servidor, número de novos cadastros e planejar pitches para patrocinadores.'],
        ['Planejamento de OKRs', 'Trimestral', 'Co-fundadores', 'Revisar as metas atingidas no trimestre anterior e reajustar metas do trimestre seguinte com dados reais.'],
    ],
    header_color='8B26FD', alt_color='F5F3FF'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 9 — REGRAS DA CASA
# ═══════════════════════════════════════════════════════════════════════════
h1(doc, 'SEÇÃO 9 — REGRAS DA CASA')
doc.add_paragraph()

regras = [
    ('REGRA 1', 'Entrega e Vesting', 'As participações societárias serão vinculadas ao cumprimento das metas individuais de entrega (desenvolvimento das telas, fechamento de parcerias, marketing).'),
    ('REGRA 2', 'Decisão de Produto', 'Definições sobre layout, features e APIs serão discutidas coletivamente, com voto técnico do programador responsável.'),
    ('REGRA 3', 'Captação e Investimento', 'Qualquer proposta de investimento anjo será analisada coletivamente, preservando a governança do grupo de fundadores.'),
    ('REGRA 4', 'Compliance de Dados', 'Os currículos cadastrados e as localizações das empresas seguirão regras básicas de privacidade e segurança da LGPD.'),
]
tbl(doc, ['#', 'Regra', 'Descrição Detalhada'], regras, header_color='1F202B', alt_color='F8F8FF')
doc.add_paragraph()

divider(doc)
doc.add_paragraph()
par(doc, 'Planejamento de OKRs e Estratégia do Projeto Vagas.me', italic=True, color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
par(doc, 'Co-piloto Antigravity  |  Itapetininga - SP  |  Julho 2026', italic=True, color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# Criar a pasta de destino caso não exista
saida_dir = r'c:\Users\WINDOWS11\Documents\Clientes\Vagas.me\saidas\planejamento'
if not os.path.exists(saida_dir):
    os.makedirs(saida_dir)

output = os.path.join(saida_dir, 'Vagas_me_Planejamento_Estrategico_OKR.docx')
doc.save(output)
print(f'OK: {output}')
